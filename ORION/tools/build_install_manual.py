#!/usr/bin/env python3
"""Genera docs/Manual_de_instalacion_ORION.docx.

Manual de instalacion completo y conciso para reproducir ORION desde cero
en una maquina limpia. Para la memoria del TFG. No incluye pasos opcionales
sin valor (capturas de progreso, comentarios obvios, ramas exploratorias).
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


OUT = Path(__file__).resolve().parent.parent / "docs" / "Manual_de_instalacion_ORION.docx"


def set_cell_shading(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def add_code(doc, code, lang_hint=None):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(code.rstrip())
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1B, 0x2A, 0x4E)
    # Sombreado gris claro al parrafo
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F4F7")
    pPr.append(shd)


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    return p


def add_bullets(doc, items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(it)


def add_h1(doc, text):
    doc.add_heading(text, level=1)


def add_h2(doc, text):
    doc.add_heading(text, level=2)


def add_h3(doc, text):
    doc.add_heading(text, level=3)


def build():
    doc = Document()

    # Estilo base
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Portada
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Manual de instalación de ORION")
    r.bold = True
    r.font.size = Pt(26)
    r.font.color.rgb = RGBColor(0x1B, 0x2A, 0x4E)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run("Optimized Real-time Inspection Of Networks")
    rs.italic = True
    rs.font.size = Pt(13)
    rs.font.color.rgb = RGBColor(0x2C, 0x52, 0x82)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(
        "Trabajo Fin de Grado en Ingeniería Informática\n"
        "Fabián Delgado Díez · Universidad Europea de Madrid\n"
        "Director: José Javier Ruiz Cobo"
    )

    doc.add_paragraph()
    intro = doc.add_paragraph()
    intro.add_run(
        "Este manual reproduce la instalación de ORION desde una máquina Ubuntu "
        "limpia hasta el sistema funcional. Se omiten pasos de exploración o "
        "configuraciones de prueba sin valor permanente. Cada comando aparece "
        "tal cual se debe ejecutar."
    )

    doc.add_page_break()

    # 1
    add_h1(doc, "1. Requisitos previos")

    add_h2(doc, "1.1. Hardware mínimo")
    add_bullets(doc, [
        "CPU x86-64 con virtualización y soporte eBPF (cualquier procesador moderno desde 2016).",
        "16 GB de RAM (Elasticsearch + Kibana + Packetbeat son los principales consumidores).",
        "100 GB de disco SSD (Elasticsearch crece rápido con la captura de tráfico).",
        "3 interfaces de red Ethernet: una de gestión y salida (1 Gb) y dos de 10 Gb para cliente y generador.",
    ])

    add_h2(doc, "1.2. Sistema operativo")
    add_bullets(doc, [
        "Ubuntu Desktop 24.04 LTS o superior.",
        "Kernel Linux 6.5 o superior (requerido por el flag BPF_FIB_LOOKUP_TBID que usa el modo XDP_REDIRECT).",
        "Acceso root o sudo sin contraseña al usuario que ejecutará ORION.",
    ])

    add_h2(doc, "1.3. Topología física")
    add_para(doc,
        "ORION es la máquina central conectada a tres redes mediante sus tres NICs. "
        "El switch o router doméstico/oficina proporciona la salida real a Internet a "
        "través de eno1; un PC cliente se conecta a enp6s0f0 (red interna emulada); el "
        "generador de tráfico TRex se conecta a enp6s0f1 cuando se ejecuten pruebas de "
        "rendimiento.")

    # 2
    add_h1(doc, "2. Paquetes base del sistema")
    add_para(doc, "Instalar las dependencias mínimas (compiladores, eBPF, redes, Python):")
    add_code(doc,
"""sudo apt update
sudo apt install -y \\
  build-essential clang llvm git curl jq \\
  linux-headers-$(uname -r) \\
  python3 python3-venv python3-pip \\
  python3-bpfcc bpfcc-tools bpftrace linux-tools-common linux-tools-generic \\
  nftables iproute2 arping isc-dhcp-client \\
  apt-transport-https gnupg""")

    # 3
    add_h1(doc, "3. Configuración de red base")

    add_h2(doc, "3.1. Migrar a systemd-networkd")
    add_para(doc,
        "ORION usa systemd-networkd en lugar de NetworkManager para tener configuración "
        "declarativa pura y evitar interferencias en las interfaces.")
    add_code(doc,
"""sudo systemctl stop NetworkManager
sudo systemctl disable NetworkManager
sudo systemctl enable --now systemd-networkd""")

    add_h2(doc, "3.2. Netplan")
    add_para(doc, "Crear el fichero /etc/netplan/01-orion.yaml con el siguiente contenido:")
    add_code(doc,
"""network:
  version: 2
  renderer: networkd
  ethernets:
    eno1:
      dhcp4: true
      dhcp6: false
    enp6s0f0:
      dhcp4: false
      dhcp6: false
      addresses:
        - 10.0.2.1/24
    enp6s0f1:
      dhcp4: false
      dhcp6: false""")
    add_para(doc, "Aplicar permisos y configuración:")
    add_code(doc,
"""sudo chmod 600 /etc/netplan/01-orion.yaml
sudo netplan apply""")

    add_h2(doc, "3.3. IP forwarding y sysctls")
    add_code(doc,
"""sudo tee /etc/sysctl.d/99-orion.conf >/dev/null <<'EOF'
net.ipv4.ip_forward=1
net.ipv4.conf.all.rp_filter=2
EOF
sudo sysctl --system""")

    # 4
    add_h1(doc, "4. Elasticsearch y Kibana nativos")

    add_h2(doc, "4.1. Repositorio Elastic 8.x")
    add_code(doc,
"""curl -fsSL https://artifacts.elastic.co/GPG-KEY-elasticsearch \\
  | sudo gpg --dearmor -o /usr/share/keyrings/elastic.gpg
echo "deb [signed-by=/usr/share/keyrings/elastic.gpg] \\
https://artifacts.elastic.co/packages/8.x/apt stable main" \\
  | sudo tee /etc/apt/sources.list.d/elastic-8.x.list
sudo apt update
sudo apt install -y elasticsearch kibana""")

    add_h2(doc, "4.2. Configuración mínima")
    add_para(doc, "Editar /etc/elasticsearch/elasticsearch.yml dejando estas opciones (todo localhost, sin seguridad para entorno de TFG):")
    add_code(doc,
"""cluster.name: orion
node.name: orion-node-1
network.host: 127.0.0.1
http.port: 9200
discovery.type: single-node
xpack.security.enabled: false
xpack.security.http.ssl.enabled: false
xpack.security.transport.ssl.enabled: false""")
    add_para(doc, "Editar /etc/kibana/kibana.yml para permitir embeber dashboards en iframes:")
    add_code(doc,
"""server.host: "127.0.0.1"
server.publicBaseUrl: "http://localhost:5601"
elasticsearch.hosts: ["http://localhost:9200"]
server.securityResponseHeaders.disableEmbedding: false
xpack.security.sameSiteCookies: "None"
xpack.security.session.idleTimeout: "24h"
csp.strict: false""")

    add_h2(doc, "4.3. Arrancar")
    add_code(doc,
"""sudo systemctl enable --now elasticsearch
sudo systemctl enable --now kibana
curl -s http://localhost:9200 | jq .""")

    # 5
    add_h1(doc, "5. Packetbeat")
    add_para(doc, "Mismo repositorio Elastic. Captura metadata de paquetes en las tres interfaces:")
    add_code(doc,
"""sudo apt install -y packetbeat""")
    add_para(doc, "Editar /etc/packetbeat/packetbeat.yml dejando las interfaces y el output a ES:")
    add_code(doc,
"""packetbeat.interfaces.devices: [eno1, enp6s0f0, enp6s0f1]
packetbeat.flows:
  enabled: true
  timeout: 30s
  period: 5s

packetbeat.protocols:
  - type: icmp
    enabled: true
  - type: dns
    ports: [53]
  - type: http
    ports: [80, 8080, 8000]
  - type: tls
    ports: [443, 8443]

output.elasticsearch:
  hosts: ["http://localhost:9200"]

setup.kibana:
  host: "http://localhost:5601" """)
    add_code(doc,
"""sudo packetbeat setup --pipelines
sudo systemctl enable --now packetbeat""")

    # 6
    add_h1(doc, "6. Clonar e instalar ORION")
    add_code(doc,
"""cd ~/Escritorio/TFG_Fabian
git clone <URL-del-repositorio> ORION
cd ORION
python3 -m venv venv
./venv/bin/pip install -r requirements.txt""")
    add_para(doc, "El directorio resultante contiene: ebpf/ (programa XDP en C), tools/ (loaders y workers Python), controlplane/ (API Flask), frontend/ (plantillas y JS), systemd/ (unidades de servicio), demo/orion_bank (banco ficticio para la demo MITM).")

    # 7
    add_h1(doc, "7. Índices de Elasticsearch")
    add_para(doc, "ORION almacena toda su configuración y telemetría en Elasticsearch. Crear los índices base:")
    add_code(doc,
"""for IDX in orion-rules orion-rules-history orion-iface-config \\
            orion-nat-mappings orion-mitm-traffic orion-alerts \\
            orion-threat-intel orion-flows orion-blocks-history; do
  curl -s -X PUT "http://localhost:9200/$IDX" -H 'Content-Type: application/json' -d'{}'
  echo
done""")
    add_para(doc, "Insertar el mapeo NAT inicial del cliente:")
    add_code(doc,
"""curl -s -X POST "http://localhost:9200/orion-nat-mappings/_doc/cliente-a" \\
  -H 'Content-Type: application/json' -d '{
    "internal_ip": "10.0.2.10",
    "external_ip": "0.0.0.0",
    "client_label": "PC cliente",
    "enabled": true
  }'""")
    add_para(doc, "El campo external_ip se sobreescribe automáticamente cuando macvlan_up.sh obtiene la IP por DHCP.")

    # 8
    add_h1(doc, "8. mitmproxy y CA")
    add_code(doc,
"""sudo apt install -y mitmproxy
mitmweb --listen-port 8888 --web-port 8081 &
sleep 5
pkill mitmweb""")
    add_para(doc, "Esto genera la CA en ~/.mitmproxy/. Hay que instalar mitmproxy-ca-cert.pem en el PC cliente como autoridad de confianza del navegador y del sistema.")

    # 9
    add_h1(doc, "9. AbuseIPDB (inteligencia de amenazas)")
    add_para(doc, "Solicitar una clave gratuita en abuseipdb.com y guardarla:")
    add_code(doc,
"""sudo mkdir -p /etc/orion
sudo tee /etc/orion/abuseipdb.env >/dev/null <<EOF
ABUSEIPDB_API_KEY=tu_clave_aqui
EOF
sudo chmod 600 /etc/orion/abuseipdb.env""")

    # 10
    add_h1(doc, "10. Servicios systemd de ORION")
    add_para(doc, "Copiar las unidades al sistema y recargar:")
    add_code(doc,
"""sudo cp systemd/*.service systemd/*.timer /etc/systemd/system/
sudo systemctl daemon-reload""")
    add_para(doc, "Habilitar y arrancar los servicios permanentes:")
    add_code(doc,
"""sudo systemctl enable --now orion-macvlan.service
sudo systemctl enable --now orion-mitmweb.service
sudo systemctl enable --now orion-waf-detector.timer
sudo systemctl enable --now orion-block-expirer.timer
sudo systemctl enable --now orion-abuseipdb.timer
sudo systemctl enable --now orion-alerts.timer""")
    add_para(doc, "Arrancar el plano de datos en uno de los dos modos (mutuamente excluyentes por Conflicts=):")
    add_code(doc,
"""# Modo Observabilidad (operativo, el habitual):
sudo systemctl enable --now orion-xdp-pass.service

# Modo Rendimiento (en desarrollo, solo si se va a benchmarkear):
# sudo systemctl enable --now orion-xdp-redirect.service
# sudo systemctl enable --now orion-flow-consumer.service""")

    # 11
    add_h1(doc, "11. Control plane (Flask)")
    add_code(doc,
"""cd ~/Escritorio/TFG_Fabian/ORION
sudo ./venv/bin/python orion.py""")
    add_para(doc, "La interfaz web queda disponible en http://localhost:5000. Para producción se puede empaquetar como servicio systemd, pero en el TFG basta con ejecutarlo en primer plano durante la sesión.")

    # 12
    add_h1(doc, "12. Configurar el PC cliente")
    add_bullets(doc, [
        "Asignar IP estática 10.0.2.10/24, gateway 10.0.2.1.",
        "DNS: 1.1.1.1 (directo) o 10.0.2.1 si se monta dnsmasq como forwarder en ORION.",
        "Instalar mitmproxy-ca-cert.pem como autoridad de confianza para que el navegador acepte los certificados firmados por la CA de ORION (paso imprescindible para descifrar HTTPS).",
    ])

    # 13
    add_h1(doc, "13. Verificación end-to-end")
    add_h2(doc, "13.1. Salud de los componentes")
    add_code(doc,
"""curl -s http://localhost:9200 | jq .version.number
sudo bpftool net show
sudo bpftool map show | grep orion
systemctl is-active elasticsearch kibana packetbeat orion-mitmweb
curl -s http://localhost:5000/api/health | jq .""")

    add_h2(doc, "13.2. Conectividad")
    add_bullets(doc, [
        "Desde ORION: ping 10.0.2.10 (cliente) y ping 8.8.8.8 (Internet propio).",
        "Desde el cliente: ping 10.0.2.1 (gateway ORION) y navegar al banco de demostración en https://10.0.3.2.",
        "Comprobar que la GUI de ORION muestra tráfico en la pestaña Resumen y los dashboards de Kibana en Dashboards.",
    ])

    # 14
    add_h1(doc, "14. Banco de demostración")
    add_para(doc, "El proyecto incluye un banco ficticio para demostrar el descifrado HTTPS y el WAF con tráfico realista propio:")
    add_code(doc,
"""cd demo/orion_bank
./run.sh""")
    add_para(doc, "Queda escuchando en https://10.0.3.2:8443. Iniciar sesión, hacer transferencias y descargar el comprobante PDF; todo aparece descifrado en ORION si MITM está activo.")

    # 15
    add_h1(doc, "15. Operación habitual")
    add_h2(doc, "15.1. Cambiar entre modos XDP")
    add_code(doc,
"""# Pasar a Rendimiento (apaga PASS automáticamente por Conflicts=):
sudo systemctl restart orion-xdp-redirect.service

# Volver a Observabilidad:
sudo systemctl restart orion-xdp-pass.service""")
    add_h2(doc, "15.2. Logs útiles")
    add_code(doc,
"""sudo journalctl -u orion-xdp-pass -f
sudo journalctl -u orion-mitmweb -f
sudo journalctl -u orion-waf-detector -f
sudo journalctl -u packetbeat -f""")
    add_h2(doc, "15.3. Inspección directa de los mapas eBPF")
    add_code(doc,
"""sudo bpftool map dump name rules
sudo bpftool map dump name nat_in2out
sudo bpftool map dump name blocked_cidrs""")

    # 16
    add_h1(doc, "16. Resolución de problemas frecuentes")
    table = doc.add_table(rows=1, cols=2)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Síntoma"
    hdr[1].text = "Causa habitual y acción"
    for cell in hdr:
        for r in cell.paragraphs[0].runs:
            r.bold = True
        set_cell_shading(cell, "2C5282")
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    rows = [
        ("Flask no arranca: módulo bcc no encontrado",
         "Instalar python3-bpfcc del sistema y arrancar Flask con sudo."),
        ("XDP no carga: verifier rechaza el programa",
         "Kernel demasiado antiguo. Actualizar a 6.5 o superior."),
        ("El cliente no sale a Internet en modo PASS",
         "Comprobar que orion-macvlan.service está activo y que la regla de policy routing existe (ip rule list)."),
        ("Cliente no navega HTTPS con MITM",
         "La CA mitmproxy-ca-cert.pem no está instalada en el cliente, o no la acepta el navegador. Reinstalar."),
        ("Packetbeat no indexa",
         "Revisar /etc/packetbeat/packetbeat.yml (interfaces correctas) y el servicio con systemctl status packetbeat."),
        ("0 eventos en orion-flows (modo REDIRECT)",
         "Mapeo NAT no coincide con la IP real del cliente. Verificar internal_ip en orion-nat-mappings."),
    ]
    for s, c in rows:
        row = table.add_row().cells
        row[0].text = s
        row[1].text = c

    # 17
    add_h1(doc, "17. Desinstalación")
    add_code(doc,
"""sudo systemctl disable --now orion-xdp-pass orion-xdp-redirect \\
  orion-flow-consumer orion-mitmweb orion-macvlan \\
  orion-waf-detector.timer orion-block-expirer.timer \\
  orion-abuseipdb.timer orion-alerts.timer
sudo rm /etc/systemd/system/orion-*.{service,timer}
sudo systemctl daemon-reload
sudo apt remove --purge elasticsearch kibana packetbeat mitmproxy
sudo rm -rf /var/lib/elasticsearch /var/lib/kibana /etc/orion""")

    doc.save(OUT)
    print(f"Generado: {OUT}")


if __name__ == "__main__":
    build()
